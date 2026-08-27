from __future__ import annotations

from pathlib import Path
from typing import Iterable

import matplotlib.pyplot as plt
from matplotlib import font_manager
from matplotlib.patches import FancyArrowPatch, FancyBboxPatch
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
PUBLIC = ROOT / "public"
FIGURES = ROOT / "artifacts" / "research-pack-figures"
OUTPUT = PUBLIC / "clinical-sociolinguistics-apa7-research-pack.docx"
KAI_PATH = Path(r"C:\Windows\Fonts\kaiu.ttf")
FONT = "DFKai-SB" if KAI_PATH.exists() else "KaiTi"

INK = "19352D"
ORANGE = "D7653E"
CREAM = "F7F3EA"
SAGE = "DCE6DE"
LINE = "C8CFC8"

URLS = {
    "chen": "https://ndltd.ncl.edu.tw/cgi-bin/gs32/gsweb.cgi?o=dnclcdr&s=id%3D%22105NCCU5100005%22.&searchmode=basic",
    "hast": "https://ndltd.ncl.edu.tw/cgi-bin/gs32/gsweb.cgi/login?o=dnclcdr&s=id%3D%22112NDHU5577003%22.&searchmode=basic",
    "lakaw": "https://ndltd.ncl.edu.tw/cgi-bin/gs32/gsweb.cgi/login?o=dnclcdr&s=id%3D%22112NDHU5577013%22.&searchmode=basic",
    "hsiao": "https://ndltd.ncl.edu.tw/cgi-bin/gs32/gsweb.cgi/login?o=dnclcdr&s=id%3D%22113NPUS5457005%22.&searchmode=basic",
    "chang": "https://ndltd.ncl.edu.tw/cgi-bin/gs32/gsweb.cgi/login?o=dnclcdr&s=id%3D%22112FJU00248043%22.&searchmode=basic",
    "chuan": "https://ndltd.ncl.edu.tw/cgi-bin/gs32/gsweb.cgi/ccd%3D4YLJ_i/search?s=id%3D%22113NKUT0836016%22.&searchmode=basic",
}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=100, bottom=90, end=100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        tag = "left" if side == "start" else "right" if side == "end" else side
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color=LINE, size=6) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:color"), color)


def set_table_widths(table, widths_cm: Iterable[float]) -> None:
    widths = list(widths_cm)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    total_twips = int(sum(widths) / 2.54 * 1440)
    tbl_w.set(qn("w:w"), str(total_twips))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width_cm in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width_cm / 2.54 * 1440)))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = Cm(widths[min(idx, len(widths) - 1)])
            cell.width = width
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width.cm / 2.54 * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def set_run_font(run, size=12, bold=False, italic=False, color=INK) -> None:
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attr}"), FONT)


def set_paragraph(p, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, before=0, after=8, line=2.0) -> None:
    p.alignment = alignment
    fmt = p.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line
    for run in p.runs:
        set_run_font(run)


def add_hyperlink(paragraph, text: str, url: str, color="0563C1") -> None:
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    for attr in ("ascii", "hAnsi", "eastAsia", "cs"):
        r_fonts.set(qn(f"w:{attr}"), FONT)
    r_pr.append(r_fonts)
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    r_pr.append(c)
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    r_pr.append(u)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "22")
    r_pr.append(sz)
    new_run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_text(doc, text: str, *, bold=False, italic=False, size=12, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, after=8, line=2.0):
    p = doc.add_paragraph()
    p.alignment = alignment
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, italic=italic)
    return p


def add_heading(doc, text: str, level=1) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    set_run_font(r, size=16 if level == 1 else 13, bold=True, color=INK if level == 1 else ORANGE)


def add_label_title(doc, kind: str, number: int, title: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(0)
    set_run_font(p.add_run(f"{kind} {number}"), bold=True, size=11)
    p2 = doc.add_paragraph()
    p2.paragraph_format.keep_with_next = True
    p2.paragraph_format.space_after = Pt(5)
    set_run_font(p2.add_run(title), italic=True, size=11)


def add_note(doc, explanation: str, links: list[tuple[str, str]]) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.25
    set_run_font(p.add_run("註。"), italic=True, size=10)
    set_run_font(p.add_run(explanation + " "), size=10)
    for i, (label, url) in enumerate(links):
        if i:
            set_run_font(p.add_run("；"), size=10)
        add_hyperlink(p, label, url)
    set_run_font(p.add_run("。"), size=10)


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[float]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_widths(table, widths)
    set_table_borders(table)
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for i, text in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_cell_shading(cell, INK)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        set_run_font(p.add_run(text), size=9, bold=True, color="FFFFFF")
    for ridx, row in enumerate(rows):
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            cell = cells[idx]
            if ridx % 2 == 1:
                set_cell_shading(cell, CREAM)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if idx == 0 else WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            set_run_font(p.add_run(text), size=8.5, bold=(idx == 0))


def configure_matplotlib() -> None:
    if KAI_PATH.exists():
        font_manager.fontManager.addfont(str(KAI_PATH))
    plt.rcParams.update({
        "font.family": FONT,
        "axes.unicode_minus": False,
        "figure.facecolor": "white",
        "savefig.facecolor": "white",
    })


def rounded(ax, xy, wh, text, face, edge=f"#{INK}", fontsize=12, text_color=f"#{INK}"):
    x, y = xy
    w, h = wh
    edge = edge if str(edge).startswith("#") else f"#{edge}"
    text_color = text_color if str(text_color).startswith("#") else f"#{text_color}"
    patch = FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.02,rounding_size=0.03", facecolor=face, edgecolor=edge, linewidth=1.4)
    ax.add_patch(patch)
    ax.text(x + w / 2, y + h / 2, text, ha="center", va="center", fontsize=fontsize, color=text_color, linespacing=1.35)
    return patch


def save_context_figure(path: Path) -> None:
    fig, ax = plt.subplots(figsize=(7.2, 3.9), layout="constrained")
    ax.set_xlim(0, 1); ax.set_ylim(0, 1); ax.axis("off")
    items = [
        (0.03, "生命史與\n制度脈絡", SAGE),
        (0.275, "社會網絡與\n共同實踐", CREAM),
        (0.52, "語言使用量、\n優勢與方言", SAGE),
        (0.765, "任務中的\n語言表現", CREAM),
    ]
    for i, (x, label, color) in enumerate(items):
        rounded(ax, (x, .55), (.19, .24), label, f"#{color}")
        if i < len(items)-1:
            ax.add_patch(FancyArrowPatch((x+.19, .67), (items[i+1][0], .67), arrowstyle="-|>", mutation_scale=16, color=f"#{ORANGE}", linewidth=1.5))
    rounded(ax, (.20, .13), (.60, .20), "調節／混淆：教育與識字、聽力、健康史、任務熟悉度、口譯者效應", "#F1DFD5", edge=ORANGE, fontsize=10.5)
    ax.add_patch(FancyArrowPatch((.50, .33), (.50, .54), arrowstyle="-|>", mutation_scale=15, color=f"#{ORANGE}", linewidth=1.3))
    ax.text(.03, .92, "概念路徑（研究假設，不代表已證實因果）", fontsize=11, color=f"#{INK}", weight="bold")
    fig.savefig(path, dpi=220, bbox_inches=None)
    plt.close(fig)


def save_dual_axis_figure(path: Path) -> None:
    fig, ax = plt.subplots(figsize=(7.2, 4.4), layout="constrained")
    ax.set_xlim(0, 10); ax.set_ylim(0, 10); ax.set_xticks([]); ax.set_yticks([])
    for spine in ax.spines.values(): spine.set_visible(False)
    ax.axvline(5, color=f"#{INK}", linewidth=1.4); ax.axhline(5, color=f"#{INK}", linewidth=1.4)
    blocks = [
        (0,5,5,5,"社會語言學高／臨床低\n語言活力、網絡、世代傳承", SAGE),
        (5,5,5,5,"社會語言學高／臨床高\n文化公平的多語臨床評估", "F1DFD5"),
        (0,0,5,5,"社會語言學低／臨床低\n描述性背景與工具盤點", CREAM),
        (5,0,5,5,"社會語言學低／臨床高\n單一測驗與診斷效度", "E5E7D9"),
    ]
    for x,y,w,h,label,color in blocks:
        ax.add_patch(FancyBboxPatch((x+.12,y+.12),w-.24,h-.24,boxstyle="round,pad=.02,rounding_size=.12",facecolor=f"#{color}",edgecolor="white"))
        ax.text(x+w/2,y+h/2,label,ha="center",va="center",fontsize=11,linespacing=1.5,color=f"#{INK}")
    ax.annotate("臨床／認知證據強度 →", xy=(.5,-.07), xycoords="axes fraction", ha="center", fontsize=10.5, color=f"#{ORANGE}")
    ax.annotate("社會語言脈絡整合度 →", xy=(-.07,.5), xycoords="axes fraction", ha="center", va="center", rotation=90, fontsize=10.5, color=f"#{ORANGE}")
    fig.savefig(path, dpi=220, bbox_inches=None)
    plt.close(fig)


def save_pipeline_figure(path: Path) -> None:
    fig, ax = plt.subplots(figsize=(7.2, 3.9), layout="constrained")
    ax.set_xlim(0,1); ax.set_ylim(0,1); ax.axis("off")
    stages = ["共同定義\n問題", "多語生命史\n與同意", "聽力／健康\n背景", "自然互動\n語料", "雙語／方言\n任務", "社群共同\n詮釋", "分級保存、\n撤回與回饋"]
    positions = [(0.02,.62),(0.27,.62),(0.52,.62),(0.77,.62),(0.145,.20),(0.395,.20),(0.645,.20)]
    for i,(label,(x,y)) in enumerate(zip(stages,positions)):
        rounded(ax,(x,y),(.20,.19),f"{i+1:02d}\n{label}","#DCE6DE" if i%2==0 else "#F7F3EA",fontsize=10)
        if i < 3:
            ax.add_patch(FancyArrowPatch((x+.20,y+.095),(positions[i+1][0],positions[i+1][1]+.095),arrowstyle="-|>",mutation_scale=14,color=f"#{ORANGE}"))
        elif i == 3:
            ax.add_patch(FancyArrowPatch((x+.10,y),(positions[i+1][0]+.10,positions[i+1][1]+.19),connectionstyle="arc3,rad=-.15",arrowstyle="-|>",mutation_scale=14,color=f"#{ORANGE}"))
        elif i < 6:
            ax.add_patch(FancyArrowPatch((x+.20,y+.095),(positions[i+1][0],positions[i+1][1]+.095),arrowstyle="-|>",mutation_scale=14,color=f"#{ORANGE}"))
    ax.text(.5,.93,"參與式、可追溯、可撤回的研究資料流程",ha="center",fontsize=11,weight="bold",color=f"#{INK}")
    fig.savefig(path,dpi=220,bbox_inches=None)
    plt.close(fig)


def make_figures() -> list[Path]:
    FIGURES.mkdir(parents=True, exist_ok=True)
    configure_matplotlib()
    paths = [FIGURES / "figure-1-context.png", FIGURES / "figure-2-dual-axis.png", FIGURES / "figure-3-pipeline.png"]
    save_context_figure(paths[0]); save_dual_axis_figure(paths[1]); save_pipeline_figure(paths[2])
    return paths


def add_page_number(section) -> None:
    p = section.header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run()
    fld_char1 = OxmlElement("w:fldChar"); fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = " PAGE "
    fld_char2 = OxmlElement("w:fldChar"); fld_char2.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char1, instr, fld_char2])
    set_run_font(run, size=10)


def build_doc(figures: list[Path]) -> None:
    PUBLIC.mkdir(parents=True, exist_ok=True)
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(1); sec.bottom_margin = Inches(1); sec.left_margin = Inches(1); sec.right_margin = Inches(1)
    add_page_number(sec)

    normal = doc.styles["Normal"]
    normal.font.name = FONT; normal.font.size = Pt(12)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.line_spacing = 2
    normal.paragraph_format.space_after = Pt(0)

    # Editorial cover
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(30); p.paragraph_format.space_after = Pt(24)
    set_run_font(p.add_run("CLINICAL SOCIOLINGUISTICS · RESEARCH PACK"), size=10, bold=True, color=ORANGE)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(18)
    set_run_font(p.add_run("《臨床社會語言學》\n圖表、研究脈絡與研究缺口整理"), size=25, bold=True)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(34)
    set_run_font(p.add_run("以高齡阿美族語、多語生命史與文化公平評估為研究轉譯焦點"), size=14, color=ORANGE)
    table = doc.add_table(rows=3, cols=2); set_table_widths(table,[3.6,11.8]); set_table_borders(table,color="FFFFFF",size=0)
    for i,(k,v) in enumerate([("文件型態","APA 第 7 版格式之研究整理包"),("字體與版面","標楷體；內文左右對齊；雙行距；四邊 1 吋"),("整理日期","2026 年 8 月 27 日")]):
        set_cell_shading(table.cell(i,0), INK); set_cell_shading(table.cell(i,1), CREAM)
        for c in table.rows[i].cells: set_cell_margins(c,160,160,160,160)
        set_run_font(table.cell(i,0).paragraphs[0].add_run(k),size=10,bold=True,color="FFFFFF")
        set_run_font(table.cell(i,1).paragraphs[0].add_run(v),size=10)
    add_text(doc,"使用聲明：本文件將閱讀心得中的概念整理為可供研究設計討論的表格與脈絡圖。圖中的箭頭表示待驗證的概念關係，不代表已建立因果效果；阿美族相關研究連結置於每一圖表的註釋中。",size=10,after=0,line=1.45)
    doc.add_page_break()

    add_heading(doc,"摘要")
    add_text(doc,"本整理包將《Clinical Sociolinguistics》（Ball, 2005）21 章閱讀心得轉譯為研究設計資源，聚焦社會網絡、多語生命史、臨床公平性、識字／語言社會化與語群個案五個主題群。文件以高齡阿美族語與神經認知研究為情境，彙整章節資料、變項操作化、概念路徑、雙軸證據框架、研究缺口優先序與參與式資料流程。每一圖表均附阿美族相關研究的可點擊連結與連接說明，以便追索證據與延伸閱讀。")
    add_heading(doc,"使用與判讀原則",2)
    for text in [
        "差異不等於障礙：方言、語碼轉換、識字經驗與文化熟悉度必須納入低分的替代解釋。",
        "生命史先於分數：同一位多語使用者的各語言能力可能隨家庭、教育、工作、遷徙與健康事件改變。",
        "圖表不宣稱因果：本文件沒有原始受試者資料，所有路徑與優先序均是研究設計假設。",
        "連結需持續核對：下列學位論文連結以臺灣博碩士論文知識加值系統為主，正式引用前宜再次檢查書目頁。",
    ]:
        p=doc.add_paragraph(style=None); p.style=doc.styles["Normal"]; p.paragraph_format.left_indent=Cm(.6); p.paragraph_format.first_line_indent=Cm(-.35)
        set_run_font(p.add_run("• "),color=ORANGE,bold=True); set_run_font(p.add_run(text)); set_paragraph(p)

    doc.add_page_break()
    add_label_title(doc,"表",1,"章節主題群、完成度與研究用途")
    add_table(doc,["主題群","章數／完成","整合重點","對阿美族研究的用途"],[
        ["社會脈絡","7／7","網絡、共同實踐、權力、文化、變遷、規劃、態度","辨識語言轉移、標準語權力與部落生活領域變化"],
        ["多語生命史","3／3","雙語動態、語碼轉換、多語習得","重建阿美語／華語的習得、使用與優勢軌跡"],
        ["臨床評估","4／6","聽力、失語、多語評估、公平性、口譯","避免將方言、聽力或文化陌生誤判為認知障礙"],
        ["識字與社會化","2／3","語言社會化、識字實踐、溝通模態","區分口語能力、羅馬字讀寫與教育經驗"],
        ["語群個案","0／2","區域／社會變異、非主流英語個案","補足『非標準不等於障礙』的比較證據"],
    ],[3.0,2.3,5.0,5.1])
    add_note(doc,"陳誼誠以語言活力指標呈現阿美語代間傳承與使用領域收縮，可支撐社會脈絡與生命史群組；Sifo Lakaw 從家庭語言去殖民化、語言意識形態與實踐社群說明家庭／部落場域的能動性。",[("陳誼誠（2017）",URLS["chen"]),("Sifo Lakaw（2024）",URLS["lakaw"])])

    add_label_title(doc,"圖",1,"從生命史與制度脈絡到任務表現的概念路徑")
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(0); p.add_run().add_picture(str(figures[0]),width=Inches(6.25))
    add_note(doc,"陳誼誠的語言活力研究提供『領域與代間傳承』背景；Hast 的數位民族誌可用來擴充數位社群／開放平台的共同實踐；本圖據此將臨床表現放回語言生命史與網絡中解讀。",[("陳誼誠（2017）",URLS["chen"]),("Hast（2023）",URLS["hast"])])

    add_label_title(doc,"表",2,"核心變項操作化矩陣")
    add_table(doc,["構念","建議指標／資料","資料型態","主要偏誤控制"],[
        ["病前語言生態","習得年齡、家庭／學校／工作語言、方言、優勢轉換","生命史訪談＋時間軸","回憶偏誤；家屬／文件交叉核對"],
        ["社會網絡","互動對象、關係強度、頻率、語言選擇、網絡縮小事件","網絡名冊＋事件紀錄","把同住誤當實際互動"],
        ["共同實踐","教會、農作、祭儀、照顧、工作與數位社群中的語言功能","參與觀察＋自然語料","研究者預設場域功能"],
        ["語言使用與優勢","每週使用比例、自評、命名／敘事、語碼轉換功能","多方法重複測量","單一分數與單語常模"],
        ["感官與健康","聽力、視力、病史、用藥、疲勞與情緒","篩檢＋病歷／自陳","將未聽清楚當成理解缺損"],
        ["評估公平性","方言、文化熟悉度、教育、識字、任務可理解性","認知訪談＋差異項目檢驗","翻譯等同文化調適"],
        ["口譯者效應","口譯者背景、提示、改述、互動回合","錄音錄影＋逐字稿標記","把口譯視為透明管道"],
        ["資料治理","同意層級、公開範圍、撤回、社群回饋、命名規則","治理協議＋稽核軌跡","一次同意涵蓋所有再利用"],
    ],[3.0,5.8,3.1,3.5])
    add_note(doc,"蕭惠美將原住民健康識能轉譯分為資訊轉譯、文化、社會結構與醫療資源等面向；全俊儒則聚焦跨文化溝通與原住民醫病關係。兩者可用於擴充『評估公平性』與『口譯者效應』欄位。",[("蕭惠美（2025）",URLS["hsiao"]),("全俊儒（2025）",URLS["chuan"])])

    add_label_title(doc,"圖",2,"臨床／認知證據與社會語言脈絡的雙軸框架")
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(0); p.add_run().add_picture(str(figures[1]),width=Inches(6.25))
    add_note(doc,"張文翊以臺東文化健康站阿美族長者進行文化敏感音樂治療，呈現文化記憶、參與和量化結果需要並置；蕭惠美的健康識能轉譯架構則提醒醫療資源與文化面向不可被單一工具取代。",[("張文翊（2024）",URLS["chang"]),("蕭惠美（2025）",URLS["hsiao"])])

    add_label_title(doc,"表",3,"研究缺口與優先序矩陣")
    add_table(doc,["優先序","缺口","可回答的問題","建議下一步"],[
        ["P1","第14、15章：非主流方言與雙語差異／障礙","低分如何由語言差異、文化陌生或神經認知變化解釋？","建立阿美語—華語雙語基線、動態評量與多重證據判讀規則"],
        ["P1","病前多語功能缺乏標準化生命史紀錄","目前表現相對於個人一生的哪一階段？","設計事件時間軸並以家屬、社群與文件交叉核對"],
        ["P1","阿美語方言／變體與任務公平性證據不足","哪些題項測到的是變體差異，而非能力？","進行認知訪談、方言審查與差異項目分析"],
        ["P2","第2、8章：社會變異比較證據未補","如何證明非標準形式不等於病理？","補讀跨語群個案並建立臨床誤判案例庫"],
        ["P2","自然互動語料與結構式測驗尚未連結","測驗表現能否代表日常溝通功能？","同日蒐集敘事、對話、任務與場域使用資料"],
        ["P2","口譯／協作者效應未被量化","提示、改述與關係位置如何改變結果？","保留雙語原始語料並標記每次協助"],
        ["P3","第20章：手語與多模態溝通未補","研究是否忽略手勢、書寫與非口語資源？","建立多模態編碼規則與可近用流程"],
        ["P3","資料主權與撤回機制欠缺實作測試","社群如何控制後續使用與公開？","共同制定分級存取、撤回演練與成果回饋週期"],
    ],[2.0,4.2,5.1,4.1])
    add_note(doc,"Sifo Lakaw 的家庭語言去殖民化研究可支持共同詮釋、語言意識形態與資料治理；Hast 對阿美語數位平台的民族誌則指向數位語料、社群參與及平台治理仍需納入。",[("Sifo Lakaw（2024）",URLS["lakaw"]),("Hast（2023）",URLS["hast"])])

    add_label_title(doc,"圖",3,"參與式且可撤回的研究資料流程")
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(0); p.add_run().add_picture(str(figures[2]),width=Inches(6.25))
    add_note(doc,"Sifo Lakaw 強調家庭與社群的語言實踐、能動性及去殖民化視角；Hast 的開放阿美語平台研究提醒數位保存不是純技術步驟，而涉及社群規範、參與和再利用。",[("Sifo Lakaw（2024）",URLS["lakaw"]),("Hast（2023）",URLS["hast"])])

    add_label_title(doc,"表",4,"阿美族相關研究的連接方式與可轉用觀點")
    add_table(doc,["研究","情境／焦點","與本整理包的連接","使用界線"],[
        ["陳誼誠（2017）","阿美族語語言活力；代間傳承與使用領域","校準病前語言生態、世代與領域變化","群體活力指標不能直接推論個人認知狀態"],
        ["Hast（2023）","阿美語萌典與 Wikipitiya 的數位民族誌","補入數位實踐社群、平台與開放語料治理","平台參與者不代表所有地區與年齡層"],
        ["Sifo Lakaw（2024）","家庭語言去殖民化、意識形態與社會化","支持共同設計、家庭語言實踐與社群詮釋","去殖民化框架不可簡化成單一介入變項"],
        ["張文翊（2024）","文化健康站阿美族長者與文化敏感音樂治療","示範文化記憶、參與及臨床量化結果並置","介入成效不能外推為語言或失智診斷效度"],
        ["蕭惠美（2025）","原住民健康識能轉譯成功因素","擴充文化、資訊轉譯、社會結構、醫療資源欄位","層級分析法的權重不是個體因果效果"],
        ["全俊儒（2025）","跨文化溝通與原住民醫病關係","支持記錄互動、信任、口譯與機構條件","單一醫院情境需經跨場域驗證"],
    ],[3.0,4.2,5.1,3.1])
    add_note(doc,"本表將六項阿美族／原住民族研究定位為『研究設計參照』，不把不同研究目的、樣本與方法合併成效果量；原文書目頁可由下列連結逐項查核。",[("語言活力",URLS["chen"]),("數位民族誌",URLS["hast"]),("家庭語言去殖民化",URLS["lakaw"]),("文化敏感音樂治療",URLS["chang"]),("健康識能轉譯",URLS["hsiao"]),("跨文化醫病溝通",URLS["chuan"])])

    add_heading(doc,"整合後的研究主張")
    add_text(doc,"若要區分高齡阿美語使用者的正常社會語言變化、語言使用減少／磨損與可能的神經認知變化，研究單位不能只是一份族語版測驗。較穩健的設計需要同時保存個人多語生命史、網絡與共同實踐、感官及健康背景、自然語料、結構式任務、口譯互動軌跡，以及社群對結果的共同詮釋。")
    add_text(doc,"最優先的證據缺口是：阿美語方言與任務公平性、病前雙語基線、自然互動與臨床任務的效標連結，以及可操作的資料主權／撤回機制。這些缺口決定研究能否把『不同』與『病理』清楚分開。")

    add_heading(doc,"參考文獻")
    refs = [
        ("Ball, M. J. (Ed.). (2005). Clinical sociolinguistics. Blackwell Publishing.", None),
        ("Hast, A. C. (2023). Mipaselak to sowal no Pangcah: A digital ethnography of two open 'Amis/Pangcah language platforms, 'Amis MoeDict and 'Amis Wikipitiya [Master's thesis, National Dong Hwa University]. ", URLS["hast"]),
        ("Sifo Lakaw. (2024). O Pangcah kami, misanoPangcah kami i loma': Misawaday a misanoholam a parod no Pangcah i Taywan [Doctoral dissertation, National Dong Hwa University]. ", URLS["lakaw"]),
        ("全俊儒（2025）。《跨文化溝通對原住民醫病關係的影響—以南投某教學醫院為例》［碩士論文，南開科技大學］。", URLS["chuan"]),
        ("張文翊（2024）。《文化敏感音樂治療應用於部落長者預防及延緩失能之成效探討》［碩士論文，輔仁大學］。", URLS["chang"]),
        ("陳誼誠（2017）。《阿美族語的語言活力》［博士論文，國立政治大學］。", URLS["chen"]),
        ("蕭惠美（2025）。《以層級分析法探討提升原住民健康識能轉譯之關鍵成功因素》［碩士論文，國立屏東科技大學］。", URLS["hsiao"]),
    ]
    for text,url in refs:
        p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(.5); p.paragraph_format.first_line_indent=Inches(-.5); p.paragraph_format.line_spacing=2; p.paragraph_format.space_after=Pt(0)
        set_run_font(p.add_run(text),size=11)
        if url: add_hyperlink(p,url,url)

    add_heading(doc,"圖表可近用描述",1)
    add_text(doc,"圖 1：四個水平方塊由左至右排列，依序為生命史與制度脈絡、社會網絡與共同實踐、語言使用量／優勢／方言、任務中的語言表現；下方列出教育識字、聽力、健康史、任務熟悉度與口譯者效應等調節或混淆因素。",size=10,line=1.5)
    add_text(doc,"圖 2：二乘二矩陣，橫軸為臨床／認知證據強度，縱軸為社會語言脈絡整合度；右上象限代表文化公平的多語臨床評估，是本研究設計的目標位置。",size=10,line=1.5)
    add_text(doc,"圖 3：七階段流程依序為共同定義問題、多語生命史與同意、聽力／健康背景、自然互動語料、雙語／方言任務、社群共同詮釋、分級保存／撤回／回饋。",size=10,line=1.5)

    props = doc.core_properties
    props.title = "《臨床社會語言學》圖表、研究脈絡與研究缺口整理"
    props.subject = "APA 7 研究整理包：高齡阿美族語與文化公平評估"
    props.author = "Clinical Sociolinguistics Reading Platform"
    props.keywords = "臨床社會語言學, 阿美族, 多語生命史, APA 7, 研究缺口"
    doc.save(OUTPUT)


if __name__ == "__main__":
    figs = make_figures()
    build_doc(figs)
    print(OUTPUT)
